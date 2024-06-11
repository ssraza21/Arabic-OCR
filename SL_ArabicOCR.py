from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from reportlab.lib.enums import TA_RIGHT
import re
import os
import pypdfium2 as pdfium
from bidi.algorithm import get_display
import arabic_reshaper
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import pytesseract
import streamlit as st
import base64
from PIL import Image, ImageOps

pdfmetrics.registerFont(TTFont('Arabic_naskh', 'NotoNaskhArabic-Regular.ttf'))
pdfmetrics.registerFont(TTFont('Urdu_naskh', 'NotoNastaliqUrdu-Regular.ttf'))
pdfmetrics.registerFont(
    TTFont('English_naskh', 'Roboto-Regular.ttf'))
language_mapping = {"Arabic": "ara", "Urdu": "urd", "English": "eng"}


def get_base64_of_bin_file(bin_file):
    with open(bin_file, 'rb') as f:
        data = f.read()
    return base64.b64encode(data).decode()


def get_img_with_href(local_img_path, target_url):
    img_format = os.path.splitext(local_img_path)[-1].replace('.', '')
    bin_str = get_base64_of_bin_file(local_img_path)
    html_code = f'''
        <a href="{target_url}">
            <img src="data:image/{img_format};base64,{bin_str}" />
        </a>'''
    return html_code


if __name__ == '__main__':
    st.markdown("# Image-to-Text Converter For Arabic and Urdu documents")
    st.markdown(
        "This application is used to convert a PDF or image that has **Arabic or Urdu** text (scanned or otherwise) and convert it into text that can be searched, copied, analyzed, etc. The output can be downloaded as a PDF or a Word Doc file for Arabic input. **Due to the weak results for Urdu, it can only be downloaded in Word doc format. You are free to make edits and export to PDF.** \n #### Please input how many pages to convert, otherwise it will default to 1. If you're unsure, click [here](https://i.imgur.com/nmTKMTc.png) for an explanation.")
    st.markdown(
        "## **Two Options** \n 1) Input a .PDF file (recommended)\n 2) Input an Image file")
    st.markdown(
        "## Notes: \n - The output for Urdu is not very good as of right now \n - Once you upload the file, a \"Running\" sign with show at the top right \n - Once it has completed running, two buttons will appear to download the respective files")
    # Import a PDF in and convert it to images. Save all of the images in a list
    st.markdown("")
    st.markdown("____________________________________")
    st.markdown(
        "### 1) Select language")
    language = st.selectbox("Language to convert", [
                            "Arabic", "Urdu", "English"])
    curr_lang = language_mapping[language]
    st.markdown(
        "### 2) Input number of pages")
    t = st.text_input("Insert only the number of pages and press enter")
    st.markdown(
        "### 3) Upload which file you'd like to convert!")
    c1, c2 = st.columns(2)
    with c1:
        pdf_input = st.file_uploader("Input a PDF")
    with c2:
        img_input = st.file_uploader("Input an Image")

    p2 = pdf_input
    images = []  # This will save all of the images that were converted
    f_name = ""
    # images = convert_from_path(pdffile)
    if pdf_input is not None or img_input is not None:

        if pdf_input is not None:
            f_name = "OCR_" + pdf_input.name[:-4]
            if t != "":
                num_of_pages = int(t)
            else:
                num_of_pages = 1
            print(num_of_pages)
            with pdfium.PdfContext(pdf_input) as pdf:
                for page in range(num_of_pages):
                    image = pdfium.render_page_topil(pdf, page)
                    images.append(image)
        else:
            f_name = "OCR_" + img_input.name[:-4]
            pil_img = Image.open(img_input)
            grey = ImageOps.grayscale(pil_img)
            images.append(grey)

        # Run Tesseract to get a list of strings where each string represents a page from the PDF
        # Lil bit less than four minutes for a 148 page PDF
        # 40 minutes for 840 page PDF!

        pages = []  # A list of strings where each string represents a page from the PDF

        print("about to edit and tesseract images")

        # For each image/page, get the text version
        for image in images:
            t = pytesseract.image_to_string(image, lang=curr_lang)
            pages.append(t)

        # init the style sheet from reportlab
        styles = getSampleStyleSheet()
        styleN = styles['BodyText']
        styleH_1 = styles['Heading1']
        styleH_3 = styles['Heading3']

        a_s_naskh = ParagraphStyle(
            'naskh', fontName=language+'_naskh', wordWrap='RTL', alignment=2, fontSize=17)
        eng_naskh = ParagraphStyle(
            'naskh', fontName=language+'_naskh', alignment=2, fontSize=17)
        style_pNum = ParagraphStyle('page_num', alignment=1, fontSize=16)

        # The list that will contain everything that will be ouputted in sequential order
        story = []
        pdf_page = 1

        # RUN THROUGH EVERY PAGE CODE
        story.append(Paragraph(
            "The Quality is NOT 100%. If there is text that is small, it will not accurately portray it", styleH_1))
        if curr_lang != "eng":
            for page in pages:
                p = "PDF page: " + str(pdf_page)

                story.append(Paragraph(p, style_pNum))
                story.append(Spacer(1, 12))
                lines = page.split("\n")
                # print(lines)
                for line in lines:
                    line = os.linesep.join([s for s in line.splitlines() if s])
                    # Run the two functions to reshape the text so it is formatted correctly for the PDF
                    print(f"line: {line}")
                    # make a paragraph of that line with the specific font + settings
                    if curr_lang == 'eng':
                        story.append(Paragraph(line, eng_naskh))
                    else:
                        reshaped_text = arabic_reshaper.reshape(line)
                        line_text = get_display(reshaped_text)
                        story.append(Paragraph(line_text, a_s_naskh))
                    story.append(Spacer(1, 12))

                # Go to the next page since we want to have it correlate to the PDF
                # story.append(PageBreak())
                pdf_page = pdf_page + 1

        # DONE WITH EVERY PAGE CODE
        story.append(
            Paragraph("Converted to digital text by Shahrukh Raza. Tool Used: https://ssraza21-arabic-ocr-sl-arabicocr-l7mosu.streamlitapp.com/ ", styleH_3))
        doc = SimpleDocTemplate(f_name+".pdf", pagesize=letter)
        doc.build(story)  # Create it with all of the information inside
        print('PDF saved Alhamdulillah')

        # Start creation of the Word Doc based on the content from tesseract
        document = Document()
        doc_style = document.styles['Normal']
        doc_font = doc_style.font
        doc_font.size = Pt(15)
        p_num = 1
        document.add_heading('OCR Reader Word Output', 0)

        for page in pages:
            filtered = filter(lambda x: not re.match(
                r'^\s*$', x), page.split("\n"))
            ppage = list(filtered)
            pp = document.add_paragraph("PDF Page: " + str(p_num) + "\n")
            for i in range(len(ppage)):
                if curr_lang == 'eng':
                    p = document.add_paragraph(ppage[i])
                else:
                    p_reshape = arabic_reshaper.reshape(ppage[i])
                    p = document.add_paragraph(p_reshape)
                    p.style.font.rtl = True
                    p.style.font.size = Pt(17)
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_num += 1
            # document.add_page_break()
        document.save(f_name+".docx")
        print("Word doc saved alhamdulillah")

        with open(f_name+".pdf", "rb") as pdf_file:
            PDFbyte = pdf_file.read()
        with open(f_name+".docx", "rb") as word_file:
            wordByte = word_file.read()
        col1, col2, col3 = st.columns(3)

        with col1:
            pass
        with col2:
            if curr_lang in ['ara']:
                st.download_button(label="Download the .PDF file", data=PDFbyte,
                                   file_name=f_name+".pdf", mime='application/octet-stream')
            st.download_button("Download the .Docx file", data=wordByte,
                               file_name=f_name+".docx", mime='application/octet-stream')
        with col3:
            pass
    st.markdown("")
    st.markdown("____________________________________")
    st.markdown(
        "##### Developed by [Shahrkuh Raza](https://ssraza21.github.io/index.html). Keep me in your Prayers InshaAllah. \n \n #### Connect with me on:")

    twitter_t = get_img_with_href(
        'Thumbnails/twitter.png', 'https://twitter.com/ssraza21')
    medium_t = get_img_with_href(
        'Thumbnails/medium.png', 'https://medium.com/@ssraza21')
    linkedin_t = get_img_with_href(
        'Thumbnails/linkedin.png', 'https://www.linkedin.com/in/ssraza21')
    github_t = get_img_with_href(
        'Thumbnails/github.png', 'https://github.com/ssraza21')

    t_1, t_2, t_3, t_4, t_5 = st.columns([1, 1, 1, 1, 4], gap="small")

    with t_1:
        st.markdown(twitter_t, unsafe_allow_html=True)
    with t_2:
        st.markdown(medium_t, unsafe_allow_html=True)
    with t_3:
        st.markdown(linkedin_t, unsafe_allow_html=True)
    with t_4:
        st.markdown(github_t, unsafe_allow_html=True)
